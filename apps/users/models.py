import datetime

from anymail.message import AnymailMessage
from django.core.urlresolvers import reverse_lazy
from django.db import models
from django.contrib.auth.models import AbstractBaseUser, BaseUserManager
from django.contrib.auth.models import Group
from django.contrib.auth.decorators import user_passes_test
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse, HttpResponseForbidden
from django.utils.translation import ugettext_lazy as _
from django.dispatch import receiver
from allauth.account.signals import user_logged_in
from auditlog.registry import auditlog
from solo.models import SingletonModel

from apps.payment.models import Payment

# imports for generating card
from django.conf import settings
from PIL import ImageFont
from PIL import Image
from PIL import ImageDraw
from urllib import urlretrieve
import os
import re

import zipfile
from StringIO import StringIO
from muscn.utils.football import get_current_season_start


class UserManager(BaseUserManager):
    def create_user(self, email, password=None, full_name=''):
        if not email:
            raise ValueError('Users must have an email address')
        user = self.model(
            email=UserManager.normalize_email(email),
            full_name=full_name,
        )
        user.set_password(password)
        user.save(using=self._db)
        return user

    def create_superuser(self, email, password=None, full_name=''):
        """
        Creates and saves a superuser with the given email, full name and password.
        """
        user = self.create_user(
            email,
            password=password,
            full_name=full_name,
        )
        user.is_superuser = True
        user.is_staff = True
        user.save(using=self._db)
        return user

    def by_group(self, group_name):
        try:
            group = Group.objects.get(name=group_name)
            return self.filter(groups=group)
        except Group.DoesNotExist:
            return False


class User(AbstractBaseUser):
    username = models.CharField(max_length=50, unique=True, blank=True, null=True)
    full_name = models.CharField(max_length=245)
    devil_no = models.PositiveIntegerField(unique=True, null=True, blank=True)
    email = models.EmailField(
        verbose_name='email address',
        max_length=254,
        unique=True,
        db_index=True)

    is_active = models.BooleanField(default=True)
    is_staff = models.BooleanField(default=False)
    is_superuser = models.BooleanField(default=False)
    groups = models.ManyToManyField(Group, related_name='users', blank=True)

    def __str__(self):
        return self.full_name or self.username or self.devil_no or self.email

    @property
    def gravatar_url(self):
        import hashlib

        return "http://www.gravatar.com/avatar/" + hashlib.md5(self.email.lower()).hexdigest() + "?s=512&d=blank"

    @property
    def profile_picture(self):
        for account in self.socialaccount_set.all():
            if account.provider == 'google' and account.extra_data.get('picture'):
                return account.extra_data.get('picture')
            if account.provider == 'facebook':
                return 'https://graph.facebook.com/' + account.uid + '/picture?type=large&width=512&height=512'
            if account.provider == 'twitter' and account.extra_data.get('screen_name'):
                return 'https://twitter.com/' + account.extra_data.get('screen_name') + '/profile_image?size=original'

        return self.gravatar_url

    @property
    def card_status(self):
        if hasattr(self, 'membership') and hasattr(self.membership, 'card_status'):
            return self.membership.get_card_status()
        return ''

    @property
    def membership_status(self):
        if self.devil_no:
            return 'Member'
        try:
            if self.membership:
                if not self.membership.payment:
                    return 'Payment information not received'
                if not self.membership.payment.verified:
                    return 'Payment not verified'
                return 'Membership not verified'
        except Membership.DoesNotExist:
            return 'Membership not applied'

    def is_member(self):
        # TODO check membership status as Approved but not Pending or Expired
        return True if hasattr(self, 'membership') and hasattr(self.membership,
                                                               'payment') and self.membership.approved_date and self.membership.approved_by and not self.membership.has_expired() else False

    USERNAME_FIELD = 'email'

    REQUIRED_FIELDS = ['username', ]

    def __unicode__(self):
        return self.full_name or self.username

    def get_short_name(self):
        # The user is identified by username
        return self.username

    def has_module_perms(self, app_label):
        "Does the user have permissions to view the app `app_label`?"
        # Simplest possible answer: Yes, always
        return True

    def has_perm(self, app_label):
        "Does the user have permissions to view the app `app_label`?"
        # Simplest possible answer: Yes, always
        return True

    def email_user(self, subject, context, text_template, html_template=None, tag='Default'):
        from django.template.loader import render_to_string

        context['user'] = self
        text_message = render_to_string(text_template, context)
        if html_template:
            html_message = render_to_string(html_template, context)
        else:
            html_message = None

        # send_mail(subject, text_message, settings.DEFAULT_FROM_EMAIL, [self.email], fail_silently=False,
        #           html_message=html_message)

        message = AnymailMessage(
            subject=subject,
            body=text_message,
            to=["%s <%s>" % (self.full_name or self.username, self.email)],
            tags=[tag],
        )
        if html_message:
            message.attach_alternative(html_message, 'text/html')
        message.track_clicks = True
        message.send()

    def is_admin(self):
        return self.is_superuser

    def in_group(self, group_name):
        try:
            group = Group.objects.get(name=group_name)
            return group in self.groups.all()
        except Group.DoesNotExist:
            return False

    def add_to_group(self, group_name):
        try:
            group = Group.objects.get(name=group_name)
            self.groups.add(group)
            return True
        except Group.DoesNotExist:
            return False

    def get_absolute_url(self):
        return reverse_lazy('view_member_profile', kwargs={'slug': self.username})

    def generate_card(self, devil_number, draw_qr, base_image):
        pk = self.pk
        name = self.full_name
        phone = self.membership.mobile

        # The co-ordinates
        devil_number_ending_xy = (882, 32)
        name_start_xy = (32, 423)
        phone_start_xy = (32, 491)
        qr_xy = (764, 298)

        # The font-sizes
        devil_number_size = 58
        name_size = 60
        phone_size = 43

        # Configuration
        # possible values: L, M, Q, H
        qr_error_correction = 'H'

        devil_number = str(devil_number)

        # Pre-process the name
        name = name.upper()
        names = name.split()
        last_name = names[-1]
        names_sans_last = names[0:-1]
        if len(names_sans_last) > 1:
            middle_names = names_sans_last[1:]
            # multiple middle names support - Amrit Bahadur Khanal Kshetri :D
            middle_name_initials_list = [middle_name[0] + '.' for middle_name in middle_names]
            middle_name_initials = '  '.join(middle_name_initials_list)
            name_sans_last = names_sans_last[0] + '  ' + middle_name_initials + ' '
        else:
            name_sans_last = names_sans_last[0] + ' '

        # Pre-process the phone number
        phone_code = None
        pattern = '^([0|\\+[0-9]{1,5})?[-\s]?([7-9][0-9]{9})$'
        matches = re.search(pattern, phone)
        if matches:
            phone_code = matches.groups()[0]
            phone = matches.groups()[1]
            if phone_code:
                phone_code = phone_code.strip('\n\t\r\+\-')
                phone_code = phone_code.lstrip('00')
            else:
                phone_code = '977'
            phone_code = '+' + phone_code + '  '

        img = Image.open(base_image)
        # img = Image.open('watermarked_with_qr.jpg')
        # img = Image.open('front.jpg')
        img = img.convert("RGBA")

        draw = ImageDraw.Draw(img)

        # write devil number
        devil_number_font = ImageFont.truetype(os.path.join(settings.STATIC_ROOT, 'fonts', 'Aileron-LightItalic.otf'),
                                               devil_number_size)
        devil_number_text_size = draw.textsize('#' + devil_number, font=devil_number_font)
        devil_number_xy = (devil_number_ending_xy[0], devil_number_ending_xy[1])
        draw.text(devil_number_xy, '#' + devil_number, fill="white", font=devil_number_font)

        # write name
        name_sans_last_font = ImageFont.truetype(os.path.join(settings.STATIC_ROOT, 'fonts', 'Aileron-Italic.otf'),
                                                 name_size)
        name_sans_last_size = draw.textsize(name_sans_last, font=name_sans_last_font)
        last_name_font = ImageFont.truetype(os.path.join(settings.STATIC_ROOT, 'fonts', 'Aileron-BoldItalic.otf'),
                                            name_size)
        last_name_size = draw.textsize(last_name, last_name_font)
        name_length = name_sans_last_size[0] + last_name_size[0]
        name_xy = (name_start_xy[0], name_start_xy[1])
        last_name_xy = (name_xy[0] + name_sans_last_size[0], name_xy[1])

        # For name Shaswotsher Adhikari having first and last name with length >= than 20
        if len(names) == 2 and len(name) >= 20:
            last_name_xy = (name_start_xy[0], name_start_xy[1])
            name_xy = (last_name_xy[0], last_name_xy[1] - 50)

        draw.text(name_xy, name_sans_last, fill="white", font=name_sans_last_font)
        draw.text(last_name_xy, last_name, fill="white", font=last_name_font)

        # write phone number
        phone_code_offset = 0
        if phone_code:
            phone_code_font = ImageFont.truetype(os.path.join(settings.STATIC_ROOT, 'fonts', 'Aileron-LightItalic.otf'),
                                                 phone_size)
            phone_code_size = draw.textsize(phone_code, phone_code_font)
            phone_code_offset = phone_code_size[0]
            phone_code_xy = (phone_start_xy[0], phone_start_xy[1])
            draw.text(phone_code_xy, phone_code, fill="white", font=phone_code_font)

        phone_font = ImageFont.truetype(os.path.join(settings.STATIC_ROOT, 'fonts', 'Aileron-Italic.otf'),
                                        phone_size)
        phone_font_size = draw.textsize(phone, phone_font)
        phone_xy = (phone_start_xy[0] + phone_code_offset, phone_start_xy[1])
        draw.text(phone_xy, phone, fill="white", font=phone_font)

        if draw_qr:
            # download qr
            if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'qrs')):
                os.makedirs(os.path.join(settings.MEDIA_ROOT, 'qrs'))
            urlretrieve(
                'http://api.qrserver.com/v1/create-qr-code/?data=https://manutd.org.np/' + devil_number + '&size=250x250&ecc=' + qr_error_correction + '&color=ffffff&bgcolor=000',
                os.path.join(settings.MEDIA_ROOT, 'qrs', str(pk) + '.png'))
            qr_path = os.path.join(settings.MEDIA_ROOT, 'qrs', str(pk) + '.png')
            qr = Image.open(qr_path)
            # make qr transparent
            qr = qr.convert('RGBA')
            data = qr.getdata()
            new_data = []
            for item in data:
                if item[0] == 0 and item[1] == 0 and item[2] == 0:
                    new_data.append((255, 255, 255, 0))
                else:
                    new_data.append(item)
            qr.putdata(new_data)
            qr.save(os.path.join(settings.MEDIA_ROOT, 'qrs', str(pk) + '.png'))
            # write qr to image
            img = img.convert('RGBA')
            img.paste(qr, qr_xy, qr)
            os.remove(qr_path)
        return img

    def get_card(self):
        base_image = os.path.join(settings.BASE_DIR, 'private', 'empty_card.jpg')
        devil_number = self.devil_no
        draw_qr = True
        img = self.generate_card(devil_number, draw_qr, base_image)
        return img

    def get_card_download(self):
        card = self.get_card()
        response = HttpResponse(content_type="image/png")
        response['Content-Disposition'] = 'attachment; filename=card_' + str(self.devil_no) + '.png'
        card.save(response, "PNG")
        return response

    def get_sample_card(self):
        from django.core.cache import cache

        cached = cache.get('sample_card_' + str(self.id))
        if cached:
            return cached

        base_image = os.path.join(settings.STATIC_ROOT, 'img', 'watermarked_with_qr.jpg')
        devil_number = u'392'
        draw_qr = False

        img = self.generate_card(devil_number, draw_qr, base_image)

        img.thumbnail((530, 325), Image.ANTIALIAS)

        if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'sample_cards')):
            os.makedirs(os.path.join(settings.MEDIA_ROOT, 'sample_cards'))

        img.save(os.path.join(settings.MEDIA_ROOT, 'sample_cards', str(self.pk) + '.png'), optimize=True)

        url = settings.MEDIA_URL + 'sample_cards/' + str(self.pk) + '.png'
        cache.set('sample_card_' + str(self.pk), url)
        return url

    def get_full_name(self):
        return self.full_name

    objects = UserManager()

    class Meta:
        ordering = ['-id']


class Membership(models.Model):
    GENDERS = (
        ('M', 'Male'),
        ('F', 'Female'),
        ('O', 'Others'),
    )
    IDENTIFICATION_TYPES = (
        ('C', 'Citizenship'),
        ('L', 'License'),
        ('I', 'Identity Card'),
    )
    MEMBERSHIP_STATUSES = (
        ('P', 'Pending'),
        ('A', 'Active'),
        ('E', 'Expired'),
    )
    SHIRT_SIZES = (
        ('S', 'S'),
        ('M', 'M'),
        ('L', 'L'),
        ('XL', 'XL'),
        ('XXL', 'XXL'),
    )
    PRESENT_STATUSES = (
        ('S', 'Student'),
        ('E', 'Employed'),
        ('U', 'Unemployed'),
    )

    user = models.OneToOneField(User, related_name='membership')
    date_of_birth = models.DateField(null=True)
    gender = models.CharField(null=True, max_length=1, choices=GENDERS)
    temporary_address = models.TextField(null=True, blank=True)
    permanent_address = models.TextField(null=True)
    homepage = models.URLField(null=True, blank=True)
    mobile = models.CharField(max_length=50, null=True)
    telephone = models.CharField(max_length=50, null=True, blank=True)
    # identification_type = models.CharField(max_length=50, null=True, choices=IDENTIFICATION_TYPES)
    identification_file = models.FileField(null=True, upload_to='identification_files/', blank=True)
    # shirt_size = models.CharField(max_length=4, null=True, choices=SHIRT_SIZES)
    # present_status = models.CharField(max_length=1, null=True, choices=PRESENT_STATUSES)
    registration_date = models.DateField(null=True, default=datetime.datetime.now)
    approved_date = models.DateField(null=True, blank=True)
    approved_by = models.ForeignKey(User, related_name='memberships_approved', null=True, blank=True)
    expiry_date = models.DateField()

    status = models.CharField(max_length=1, choices=MEMBERSHIP_STATUSES, null=True)
    payment = models.ForeignKey(Payment, blank=True, null=True, related_name='payment_for', on_delete=models.SET_NULL)

    def save(self, *args, **kwargs):
        if not self.id and not self.expiry_date:
            self.expiry_date = get_current_season_start() + datetime.timedelta(days=365)
        if not self.registration_date:
            self.registration_date = datetime.datetime.now()
        if not self.status:
            self.status = 'P'
        return super(Membership, self).save(*args, **kwargs)

    def has_expired(self):
        if self.expiry_date < datetime.date.today():
            return True
        return False

    def approved(self):
        return True if self.payment and self.payment.verified and self.approved_by else False

    def approvable(self):
        return True if self.payment and self.payment.verified else False

    def get_card_status(self):
        return self.card_status.get_status()

    def set_card_status(self, num):
        card_status, created = CardStatus.objects.get_or_create(membership=self)
        if not card_status.status == num:
            card_status.status = num
            card_status.save()

    def get_absolute_url(self):
        return reverse_lazy('update_membership', kwargs={'pk': self.pk})

    def __unicode__(self):
        return unicode(self.user)

    class Meta:
        ordering = ['-id']


class Renewal(models.Model):
    membership = models.ForeignKey(Membership, related_name='renewals')
    payment = models.ForeignKey(Payment, related_name='renewals')
    date = models.DateField(default=datetime.date.today)
    approved_date = models.DateField(null=True, blank=True)
    approved_by = models.ForeignKey(User, related_name='renewals_approved', null=True, blank=True)

    def __unicode__(self):
        return unicode(self.membership.user)


class CardStatus(models.Model):
    membership = models.OneToOneField(Membership, related_name='card_status')
    STATUSES = (
        (1, 'Awaiting Print'),
        (2, 'Printed'),
        (3, 'Delivered'),
    )
    status = models.PositiveIntegerField(choices=STATUSES, default=1)
    remarks = models.CharField(max_length=255, null=True, blank=True)

    def get_status(self):
        ret = self.get_status_display()
        if self.remarks:
            ret += ' [' + self.remarks + ']'
        return ret

    def notify(self):
        if self.status == 1:
            subject = 'Your MUSCN membership has been approved.'
            params = {}
            text_template = 'users/email/status_awaiting.txt'
            # html_template = 'users/email/status_awaiting.html'
        elif self.status == 2:
            subject = 'Your MUSCN membership card has been printed.'
            params = {}
            text_template = 'users/email/status_printed.txt'
        elif self.status == 3:
            subject = 'Your MUSCN membership card has been picked up.'
            params = {}
            text_template = 'users/email/status_delivered.txt'
        self.membership.user.email_user(subject, params, text_template)

    def __unicode__(self):
        return self.membership.user.full_name + ' - ' + self.get_status()

    class Meta:
        verbose_name_plural = 'Card Statuses'


class StaffOnlyMixin(object):
    def dispatch(self, request, *args, **kwargs):
        u = request.user
        if u.is_authenticated():
            # if bool(u.groups.filter(name__in=group_names)) | u.is_superuser():
            # return True
            if bool(u.groups.filter(name='Staff')):
                return super(StaffOnlyMixin, self).dispatch(request, *args, **kwargs)
        return HttpResponseForbidden()


class MembershipSetting(SingletonModel):
    open = models.BooleanField(default=True)
    membership_fee = models.FloatField(verbose_name='Membership Fee', blank=True, null=True)
    enable_esewa = models.BooleanField(default=True)
    welcome_letter_content = models.TextField(blank=True, null=True)

    def __str__(self):
        return 'Membership Settings'


def group_required(group, login_url=None, raise_exception=True):
    def check_perms(user):
        if user.is_authenticated:
            if isinstance(group, (str,)):
                groups = (group,)
            else:
                groups = group
            if user.groups.filter(name__in=groups).exists():
                return True
            if raise_exception:
                raise PermissionDenied
        return False

    return user_passes_test(check_perms, login_url=login_url)


class GroupProxy(Group):
    class Meta:
        proxy = True
        verbose_name = _('Group')
        # verbose_name_plural = _('Groups')


# @receiver(user_signed_up)
@receiver(user_logged_in)
def get_extra_data(request, user, sociallogin=None, **kwargs):
    if sociallogin:
        extra_data = sociallogin.account.extra_data
        if sociallogin.account.provider == 'twitter':
            user.full_name = extra_data['name']

        if sociallogin.account.provider == 'facebook':
            user.full_name = extra_data['name']
            if extra_data['gender'] == 'male':
                user.gender = 'M'
            elif extra_data['gender'] == 'female':
                user.gender = 'F'

        if sociallogin.account.provider == 'google':
            pass
            # user.first_name = sociallogin.account.extra_data['given_name']
            # user.last_name = sociallogin.account.extra_data['family_name']
            # verified = sociallogin.account.extra_data['verified_email']
            # picture_url = sociallogin.account.extra_data['picture']

        user.save()


auditlog.register(Membership)
auditlog.register(Renewal)
auditlog.register(CardStatus)


def get_members_summary():
    from docx import Document

    document = Document()
    document.add_page_break()
    members = Membership.objects.filter(user__devil_no__isnull=False)
    for member in members:
        p = document.add_paragraph('')
        p.add_run('#' + str(member.user.devil_no))
        p.add_run('\n')
        p.add_run(member.user.full_name)
        p.add_run('\n')
        p.add_run(member.mobile)
        document.add_page_break()
    document.save('/tmp/members.docx')


def initialize_card_statuses():
    memberships = Membership.objects.all()
    for membership in memberships:
        if hasattr(membership, 'card_status') or not membership.user.devil_no:
            continue
        card_status = CardStatus(membership=membership, status=3)
        card_status.save()


def get_new_cards():
    in_memory_file = StringIO()
    awaiting_cards = CardStatus.objects.filter(status=1).select_related('membership__user')
    zip_file = zipfile.ZipFile(in_memory_file, 'w')
    min = float('inf')
    max = 0
    for awaiting_card in awaiting_cards:
        if awaiting_card.membership.user.devil_no:
            fake_file = StringIO()
            devil_no = awaiting_card.membership.user.devil_no
            if devil_no < min:
                min = devil_no
            if devil_no > max:
                max = devil_no
            card = awaiting_card.membership.user.get_card()
            card.save(fake_file, "jpeg")
            zip_file.writestr(str(devil_no) + '.jpg', fake_file.getvalue())
    zip_file.close()
    name = str(min) + '-' + str(max) + '.zip'
    return name, in_memory_file


# def get_birthday_users():
#     from njango import nepdate
#     from django.db.models import Q
# 
#     today = datetime.date.today()
#     ad_month = today.strftime("%m")
#     ad_day = today.strftime("%d")
#     bs_month = nepdate.today()[1]
#     bs_day = nepdate.today()[2]
# 
#     curr_year = today.year
#     min_year_ad = curr_year - 51
#     max_year_ad = curr_year - 14
#     min_year_bs = curr_year + 7
#     max_year_bs = curr_year + 43
# 
#     return User.objects.filter(Q(membership__date_of_birth__month=bs_month, membership__date_of_birth__day=bs_day,
#                                  membership__date_of_birth__year__gte=min_year_bs,
#                                  membership__date_of_birth__year__lte=max_year_bs) | Q(
#         membership__date_of_birth__month=ad_month, membership__date_of_birth__day=ad_day,
#         membership__date_of_birth__year__gte=min_year_ad, membership__date_of_birth__year__lte=max_year_bs))


def get_birthday_users():
    from njango import nepdate
    from django.db.models import Q

    ad_year = datetime.date.today().year
    ad_month = datetime.date.today().month
    ad_day = datetime.date.today().day
    bs_year = nepdate.today()[0]
    bs_month = nepdate.today()[1]
    bs_day = nepdate.today()[2]

    users = User.objects.filter(Q(membership__date_of_birth__month=bs_month, membership__date_of_birth__day=bs_day) | Q(
        membership__date_of_birth__month=ad_month, membership__date_of_birth__day=ad_day)).select_related('membership')
    real_birthday_users = []
    for user in users:
        dob_month = user.membership.date_of_birth.month
        dob_year = user.membership.date_of_birth.year
        if dob_month == ad_month and ad_year - 40 <= dob_year <= ad_year - 10:
            real_birthday_users.append(user)
        if dob_month == bs_month and bs_year - 40 <= dob_year <= bs_year - 10:
            real_birthday_users.append(user)
    return real_birthday_users


def email_birthday_users():
    params = {}
    text_template = 'users/email/happy_birthday.txt'
    users = get_birthday_users()
    for user in users:
        subject = 'Happy birthday, ' + user.full_name.split()[0] + '.'
        user.email_user(subject, params, text_template, tag='Birthday')
